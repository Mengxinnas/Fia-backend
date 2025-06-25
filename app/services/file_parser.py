import os
import re
from typing import Optional

class FileParserService:
    """文件解析服务"""
    
    async def extract_text(self, file_path: str, file_ext: str) -> str:
        """根据文件类型提取文本"""
        try:
            if file_ext == '.txt':
                return await self._extract_text_from_txt(file_path)
            elif file_ext == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return await self._extract_text_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return await self._extract_text_from_excel(file_path)
            elif file_ext == '.csv':
                return await self._extract_text_from_csv(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {file_ext}")
        except Exception as e:
            print(f"文件解析错误: {str(e)}")
            raise
    
    async def _extract_text_from_txt(self, file_path: str) -> str:
        """提取TXT文件文本"""
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"成功使用 {encoding} 编码读取文件")
                return self._clean_text(text)
            except UnicodeDecodeError:
                continue
        raise ValueError("无法解码文件内容")
    
    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """提取PDF文件文本"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            all_text = []
            
            for page_num, page in enumerate(doc):
                # 直接提取文本
                text = page.get_text()
                if text.strip():
                    all_text.append(f"第{page_num + 1}页:\n{text.strip()}")
                
                # 尝试提取表格
                try:
                    tables = page.find_tables()
                    for table in tables:
                        table_content = table.extract()
                        table_text = []
                        for row in table_content:
                            if row and any(cell for cell in row):
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                table_text.append(row_text)
                        if table_text:
                            all_text.append("表格内容:\n" + "\n".join(table_text))
                except Exception as e:
                    print(f"表格提取错误: {e}")
            
            doc.close()
            return self._clean_text("\n\n".join(all_text))
            
        except ImportError:
            raise ValueError("PDF处理需要安装PyMuPDF库")
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")
    
    async def _extract_text_from_docx(self, file_path: str) -> str:
        """提取DOCX文件文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            all_text = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        all_text.append(" | ".join(row_text))
            
            return self._clean_text("\n".join(all_text))
            
        except ImportError:
            raise ValueError("DOCX处理需要安装python-docx库")
        except Exception as e:
            raise ValueError(f"DOCX解析失败: {str(e)}")
    
    async def _extract_text_from_excel(self, file_path: str) -> str:
        """提取Excel文件文本"""
        try:
            import pandas as pd
            
            # 读取所有工作表
            xl_file = pd.ExcelFile(file_path)
            all_text = []
            
            for sheet_name in xl_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # 添加工作表名称
                all_text.append(f"工作表: {sheet_name}")
                
                # 添加列名
                if not df.empty:
                    headers = " | ".join(str(col) for col in df.columns)
                    all_text.append(f"列名: {headers}")
                    
                    # 添加数据行
                    for _, row in df.iterrows():
                        row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row)
                        if row_text.strip():
                            all_text.append(row_text)
                
                all_text.append("")  # 工作表之间添加空行
            
            return self._clean_text("\n".join(all_text))
            
        except ImportError:
            raise ValueError("Excel处理需要安装pandas和openpyxl库")
        except Exception as e:
            raise ValueError(f"Excel解析失败: {str(e)}")
    
    async def _extract_text_from_csv(self, file_path: str) -> str:
        """提取CSV文件文本"""
        try:
            import pandas as pd
            
            # 尝试不同编码
            for encoding in ['utf-8', 'gbk', 'gb2312']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("无法解码CSV文件")
            
            all_text = []
            
            # 添加列名
            if not df.empty:
                headers = " | ".join(str(col) for col in df.columns)
                all_text.append(f"列名: {headers}")
                
                # 添加数据行
                for _, row in df.iterrows():
                    row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row)
                    if row_text.strip():
                        all_text.append(row_text)
            
            return self._clean_text("\n".join(all_text))
            
        except ImportError:
            raise ValueError("CSV处理需要安装pandas库")
        except Exception as e:
            raise ValueError(f"CSV解析失败: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        if not text:
            return ""
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符但保留中文、英文、数字和基本符号
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,\-()（）\[\]【】{}""''：:;；!！?？]', '', text)
        
        # 移除连续的标点符号
        text = re.sub(r'[.,\-]{3,}', '...', text)
        
        return text.strip() 